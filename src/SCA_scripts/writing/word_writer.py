from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt

# def write_to_docx(str, file_name = None)-> None:
#     """
#     Write the given string to a word document

#     :param str: The string to write
#     :file_name: Name of the output file. (Default: 1)
#     """
#     doc = Document()
#     p = doc.add_paragraph()
#     r = p.add_run()
#     r.text = str
#     r.font.size = Pt(12)
#     r.font.name = "David"
#     r.font.rtl = True
#     if not file_name:
#         file_name = "1"
#     file_name = file_name+".docx"
#     doc.save(file_name)
#     return

def write_to_docx(text, file_name=None) -> None:
    """
    Write the given string to a Word document (RTL, font=David)
    """
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(text)
    
    # Set font size and name
    run.font.size = Pt(12)
    run.font.name = "David"

    # For complex script fonts (like Hebrew/Arabic)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "David")
    rFonts.set(qn("w:hAnsi"), "David")
    rFonts.set(qn("w:cs"), "David")

    # Set paragraph direction to Right-to-Left
    pPr = p._element.get_or_add_pPr()
    bidi = parse_xml(r'<w:bidi w:val="1" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    pPr.append(bidi)

    if not file_name:
        file_name = "1"
    if not file_name.endswith(".docx"):
        file_name += ".docx"
    
    doc.save(file_name)
